import docx
from docx.shared import Inches
import os

def embed_diagrams_in_draft2(doc_path):
    doc = docx.Document(doc_path)
    
    elec_img_path = "final report/images/electrical_architecture.png"
    flow_img_path = "final report/images/system_flowchart.png"
    
    added_elec = False
    added_flow = False
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        
        # Insert electrical architecture image after paragraph 213 (or text matching Figure 3.3.1.3.1)
        if "Figure 3.3.1.3.1:" in txt and not added_elec:
            # Update text caption for clarity
            p.text = "Figure 3.3.1.3.1: AEGIS System Electrical Architecture showing interconnectivity between Overhead Camera, NVIDIA Jetson AGX Orin, ESP32 Sensor Hub, Load Cell Platforms, Buzzer Alert, and HMI Display."
            p.style = 'Normal'
            # Insert image in the paragraph right above the caption
            p_img = doc.paragraphs[i-1] if i > 0 else p
            run = p_img.add_run()
            run.add_picture(elec_img_path, width=Inches(6.0))
            added_elec = True
            print("Successfully embedded electrical architecture image.")
            
        # Insert system flowchart image after paragraph 246 (or text matching Figure 3.4.2.2)
        if "Figure 3.4.2.2:" in txt and not added_flow:
            p.text = "Figure 3.4.2.2: System Flowchart illustrating concurrent thread processing (CV Thread, ESP32 Load Cell Thread, HMI Server) writing to thread-safe shared state and evaluating FSM decision logic."
            p.style = 'Normal'
            p_img = doc.paragraphs[i-1] if i > 0 else p
            run = p_img.add_run()
            run.add_picture(flow_img_path, width=Inches(5.5))
            added_flow = True
            print("Successfully embedded system flowchart image.")
            
    doc.save(doc_path)
    print(f"Saved doc with embedded diagrams to {doc_path}")

if __name__ == '__main__':
    embed_diagrams_in_draft2("final report/draft2.docx")
