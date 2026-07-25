import docx
import os

def update_chapter_3(doc_path, out_path, ch3_text_path):
    if not os.path.exists(doc_path):
        print(f"File {doc_path} does not exist.")
        return

    doc = docx.Document(doc_path)
    
    ch3_idx = -1
    ch4_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        if 'chapter 3: prototype development' in text or 'chapter 3: prototyping' in text or 'chapter 3' in text:
            if ch3_idx == -1 and ('chapter 3' in text):
                ch3_idx = i
        if 'chapter 4' in text:
            if ch3_idx != -1 and ch4_idx == -1 and i > ch3_idx:
                ch4_idx = i
                break
            
    if ch3_idx != -1 and ch4_idx != -1:
        print(f"Found Ch3 at {ch3_idx}, Ch4 at {ch4_idx} in {doc_path}")
        
        # Delete paragraphs between Ch3 and Ch4
        for _ in range(ch4_idx - ch3_idx - 1):
            p = doc.paragraphs[ch3_idx + 1]
            p._element.getparent().remove(p._element)
            
        # Set Chapter 3 Heading
        doc.paragraphs[ch3_idx].text = "Chapter 3: Prototype Development"
        doc.paragraphs[ch3_idx].style = 'Heading 1'
            
        # Read the updated text content
        with open(ch3_text_path, "r") as f:
            content = f.read()
            
        ch4_p = doc.paragraphs[ch3_idx + 1]
        
        paragraphs = content.split('\n\n')
        # Skip header lines if present
        for block in paragraphs:
            block_str = block.strip()
            if not block_str or block_str.startswith('====') or block_str.startswith('CHAPTER 3'):
                continue
                
            new_p = ch4_p.insert_paragraph_before()
            if block_str.startswith('3.1 ') or block_str.startswith('3.2 ') or block_str.startswith('3.3 '):
                new_p.style = 'Heading 2'
                new_p.text = block_str
            elif block_str.startswith('3.1.') or block_str.startswith('3.2.') or block_str.startswith('3.3.'):
                new_p.style = 'Heading 3'
                new_p.text = block_str
            elif block_str.startswith('Figure ') or block_str.startswith('Table '):
                new_p.style = 'Normal'
                run = new_p.add_run(block_str)
                run.bold = True
            else:
                new_p.style = 'Normal'
                new_p.text = block_str
        
        doc.save(out_path)
        print("Successfully updated doc saved to", out_path)
    else:
        print(f"Could not find Chapter 3 or 4 bounds in {doc_path}. Found Ch3: {ch3_idx}, Ch4: {ch4_idx}")

if __name__ == '__main__':
    ch3_source = "/home/jw/CDE3301/AEGIS/final report/IS305_Final_Report_Text.txt"
    update_chapter_3("/home/jw/CDE3301/AEGIS/final report/draft2.docx", "/home/jw/CDE3301/AEGIS/final report/draft2.docx", ch3_source)
    update_chapter_3("/home/jw/CDE3301/AEGIS/final report/IS305 - Final Report_draft2_filled.docx", "/home/jw/CDE3301/AEGIS/final report/IS305 - Final Report_draft2_filled.docx", ch3_source)
