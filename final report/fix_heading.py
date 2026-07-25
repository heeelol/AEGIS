import docx

def fix_heading(doc_path):
    doc = docx.Document(doc_path)
    
    # Insert Heading 3 for 3.2.2 Load Cell System right before 3.2.2.1 Requirements if missing
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "3.2.2.1 Requirements":
            prev_p = doc.paragraphs[i-1]
            if "3.2.2 Load Cell System" not in prev_p.text:
                new_p = p.insert_paragraph_before("3.2.2 Load Cell System")
                new_p.style = 'Heading 3'
                print("Inserted Heading 3: 3.2.2 Load Cell System")
                break
                
    doc.save(doc_path)

if __name__ == '__main__':
    fix_heading("final report/draft2.docx")
